# -*- coding: utf-8 -*-
"""
Build or reformat a DOCX deliverable from text, Markdown, Word-like inputs,
or another DOCX style reference.

Pipeline:
1. Convert source content to DOCX when needed.
2. Apply a built-in, custom, or reference-DOCX style preset.
3. Post-process styles, page geometry, tables, footer page numbers,
   and optional reference-list indentation.

Run from PowerShell:
  & "<python>" "<word-skill>\\scripts\\format_docx_document.py" --check-env --source "<source-file>"
"""
from __future__ import annotations

import argparse
from copy import deepcopy
import json
import os
import posixpath
import re
import shutil
import subprocess
import sys
from datetime import datetime
from pathlib import Path
from zipfile import BadZipFile, ZipFile

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import RGBColor
from docx.shared import Cm, Pt


SCRIPT_DIR = Path(__file__).resolve().parent
SKILL_DIR = SCRIPT_DIR.parent
ASSET_REFERENCE_DOCX = SKILL_DIR / "assets" / "reference.docx"
CUSTOM_PRESETS_DIR = SKILL_DIR / "assets" / "custom_presets"
DEFAULT_BUILD_DIR_NAME = "_docx_style_build"
DEFAULT_OUTPUT_DIR_NAME = "docx_output"
DEFAULT_DOCUMENT_TITLE = "文档标题"
TITLE_STYLE_NAME = "Document Title"
TEXT_INPUT_EXTENSIONS = {".md", ".markdown", ".mdown", ".txt"}
DOCX_INPUT_EXTENSIONS = {".docx"}
LIBREOFFICE_INPUT_EXTENSIONS = {".doc", ".rtf", ".odt"}
CITATION_PATTERN = re.compile(r"\[(?:[1-9]\d*)(?:\s*(?:[-\u2010-\u2015~～,，、]\s*[1-9]\d*))*\]")

BUILTIN_PRESETS = {
    "general": {
        "page": {"width_cm": 21.0, "height_cm": 29.7, "top_cm": 2.54, "bottom_cm": 2.54, "left_cm": 2.54, "right_cm": 2.54},
        "fonts": {"east_asia": "宋体", "ascii": "Times New Roman"},
        "body": {"size": 10.5, "line": 1.35, "first_indent_cm": 0.74, "space_after": 0},
        "title": {"size": 18, "bold": True, "line": 1.2, "space_after": 14},
        "headings": {"h1": 14, "h2": 12, "h3": 10.5},
    },
    "academic-paper": {
        "page": {"width_cm": 21.0, "height_cm": 29.7, "top_cm": 2.0, "bottom_cm": 2.5, "left_cm": 2.5, "right_cm": 2.5},
        "fonts": {"east_asia": "宋体", "ascii": "Times New Roman"},
        "body": {"size": 10.5, "line": 1.5, "first_indent_cm": 0.74, "space_after": 0},
        "title": {"size": 22, "bold": True, "line": 1.2, "space_after": 18},
        "headings": {"h1": 14, "h2": 10.5, "h3": 10.5},
    },
    "report": {
        "page": {"width_cm": 21.0, "height_cm": 29.7, "top_cm": 2.7, "bottom_cm": 2.5, "left_cm": 2.8, "right_cm": 2.6},
        "fonts": {"east_asia": "宋体", "ascii": "Times New Roman"},
        "body": {"size": 11, "line": 1.45, "first_indent_cm": 0.78, "space_after": 2},
        "title": {"size": 20, "bold": True, "line": 1.2, "space_after": 16},
        "headings": {"h1": 15, "h2": 12, "h3": 11},
    },
    "briefing": {
        "page": {"width_cm": 21.0, "height_cm": 29.7, "top_cm": 2.2, "bottom_cm": 2.2, "left_cm": 2.4, "right_cm": 2.4},
        "fonts": {"east_asia": "微软雅黑", "ascii": "Arial"},
        "body": {"size": 10.5, "line": 1.28, "first_indent_cm": 0.0, "space_after": 4},
        "title": {"size": 18, "bold": True, "line": 1.15, "space_after": 12},
        "headings": {"h1": 14, "h2": 12, "h3": 10.5},
    },
    "speech": {
        "page": {"width_cm": 21.0, "height_cm": 29.7, "top_cm": 2.5, "bottom_cm": 2.5, "left_cm": 2.7, "right_cm": 2.7},
        "fonts": {"east_asia": "宋体", "ascii": "Times New Roman"},
        "body": {"size": 12, "line": 1.6, "first_indent_cm": 0.85, "space_after": 2},
        "title": {"size": 20, "bold": True, "line": 1.25, "space_after": 18},
        "headings": {"h1": 15, "h2": 13, "h3": 12},
    },
    "reference": {
        "page": {"width_cm": 21.0, "height_cm": 29.7, "top_cm": 2.54, "bottom_cm": 2.54, "left_cm": 2.54, "right_cm": 2.54},
        "fonts": {"east_asia": "宋体", "ascii": "Times New Roman"},
        "body": {"size": 10.5, "line": 1.35, "first_indent_cm": 0.74, "space_after": 0},
        "title": {"size": 18, "bold": True, "line": 1.2, "space_after": 14},
        "headings": {"h1": 14, "h2": 12, "h3": 10.5},
    },
    "english-general": {
        "page": {"width_cm": 21.0, "height_cm": 29.7, "top_cm": 2.54, "bottom_cm": 2.54, "left_cm": 2.54, "right_cm": 2.54},
        "fonts": {"east_asia": "Times New Roman", "ascii": "Times New Roman"},
        "body": {"size": 11, "line": 1.15, "first_indent_cm": 0.0, "space_after": 6},
        "title": {"size": 20, "bold": True, "line": 1.15, "space_after": 14},
        "headings": {"h1": 15, "h2": 13, "h3": 11},
    },
    "english-academic": {
        "page": {"width_cm": 21.0, "height_cm": 29.7, "top_cm": 2.54, "bottom_cm": 2.54, "left_cm": 2.54, "right_cm": 2.54},
        "fonts": {"east_asia": "Times New Roman", "ascii": "Times New Roman"},
        "body": {"size": 12, "line": 1.5, "first_indent_cm": 1.27, "space_after": 0},
        "title": {"size": 18, "bold": True, "line": 1.2, "space_after": 18},
        "headings": {"h1": 14, "h2": 12, "h3": 12},
    },
    "english-report": {
        "page": {"width_cm": 21.0, "height_cm": 29.7, "top_cm": 2.54, "bottom_cm": 2.54, "left_cm": 2.54, "right_cm": 2.54},
        "fonts": {"east_asia": "Arial", "ascii": "Arial"},
        "body": {"size": 11, "line": 1.25, "first_indent_cm": 0.0, "space_after": 6},
        "title": {"size": 20, "bold": True, "line": 1.15, "space_after": 14},
        "headings": {"h1": 15, "h2": 13, "h3": 11},
    },
    "english-briefing": {
        "page": {"width_cm": 21.0, "height_cm": 29.7, "top_cm": 2.2, "bottom_cm": 2.2, "left_cm": 2.4, "right_cm": 2.4},
        "fonts": {"east_asia": "Arial", "ascii": "Arial"},
        "body": {"size": 10.5, "line": 1.2, "first_indent_cm": 0.0, "space_after": 5},
        "title": {"size": 18, "bold": True, "line": 1.15, "space_after": 12},
        "headings": {"h1": 14, "h2": 12, "h3": 10.5},
    },
    "english-speech": {
        "page": {"width_cm": 21.0, "height_cm": 29.7, "top_cm": 2.54, "bottom_cm": 2.54, "left_cm": 2.54, "right_cm": 2.54},
        "fonts": {"east_asia": "Times New Roman", "ascii": "Times New Roman"},
        "body": {"size": 12, "line": 1.5, "first_indent_cm": 0.0, "space_after": 6},
        "title": {"size": 20, "bold": True, "line": 1.2, "space_after": 16},
        "headings": {"h1": 15, "h2": 13, "h3": 12},
    },
}


def force_black_color(r_pr):
    color = r_pr.find(qn("w:color"))
    if color is None:
        color = OxmlElement("w:color")
        r_pr.append(color)
    color.set(qn("w:val"), "000000")
    for attr in ("themeColor", "themeTint", "themeShade"):
        color.attrib.pop(qn(f"w:{attr}"), None)


def read_text_compat(path: Path) -> str:
    last_error: UnicodeDecodeError | None = None
    for encoding in ("utf-8-sig", "utf-8", "gb18030"):
        try:
            return path.read_text(encoding=encoding)
        except UnicodeDecodeError as exc:
            last_error = exc
    if last_error:
        raise last_error
    return path.read_text()


def set_run_font(run, east_asia: str = "宋体", ascii_font: str = "Times New Roman", size_pt: float | None = None, bold: bool | None = None):
    run.font.name = ascii_font
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    r_fonts.set(qn("w:eastAsia"), east_asia)
    r_fonts.set(qn("w:ascii"), ascii_font)
    r_fonts.set(qn("w:hAnsi"), ascii_font)
    if size_pt is not None:
        run.font.size = Pt(size_pt)
    if bold is not None:
        run.bold = bold
    run.font.color.rgb = RGBColor(0, 0, 0)
    force_black_color(r_pr)


def set_style_font(style, east_asia: str, ascii_font: str, size_pt: float, bold: bool = False):
    style.font.name = ascii_font
    style.font.size = Pt(size_pt)
    style.font.bold = bold
    style.font.color.rgb = RGBColor(0, 0, 0)
    r_pr = style._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    r_fonts.set(qn("w:eastAsia"), east_asia)
    r_fonts.set(qn("w:ascii"), ascii_font)
    r_fonts.set(qn("w:hAnsi"), ascii_font)
    force_black_color(r_pr)


def remove_paragraph_borders(paragraph_or_style):
    element = paragraph_or_style._element
    p_pr = element.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is not None:
        p_pr.remove(p_bdr)


def remove_paragraph_style(paragraph):
    p_pr = paragraph._p.get_or_add_pPr()
    p_style = p_pr.find(qn("w:pStyle"))
    if p_style is not None:
        p_pr.remove(p_style)


def get_or_create_paragraph_style(doc: Document, name: str):
    try:
        return doc.styles[name]
    except KeyError:
        return doc.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)


def set_cell_margins(cell, top=80, start=80, bottom=80, end=80):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_shading(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.first_child_found_in("w:shd")
    if shading is None:
        shading = OxmlElement("w:shd")
        tc_pr.append(shading)
    shading.set(qn("w:fill"), fill)
    shading.set(qn("w:val"), "clear")


def display_width(text: str) -> int:
    return sum(2 if ord(ch) > 127 else 1 for ch in text)


def is_short_value(text: str) -> bool:
    stripped = text.strip()
    return bool(re.fullmatch(r"[\d.●○◐%/\\\-\s]+", stripped)) or display_width(stripped) <= 6


def preferred_column_widths(table, total_width: int = 9070) -> list[int]:
    col_count = len(table.columns)
    if col_count <= 0:
        return []
    header = [cell.text.strip() for cell in table.rows[0].cells] if table.rows else [""] * col_count

    if col_count == 2:
        widths = [int(total_width * 0.35), total_width - int(total_width * 0.35)]
    elif col_count == 3 and header and "序号" in header[0]:
        widths = [850, int(total_width * 0.34), total_width - 850 - int(total_width * 0.34)]
    elif col_count >= 8:
        first = 1850
        rest = max(620, int((total_width - first) / (col_count - 1)))
        widths = [first] + [rest] * (col_count - 1)
    else:
        weights: list[float] = []
        for idx in range(col_count):
            samples = [row.cells[idx].text.strip() for row in table.rows[:20] if idx < len(row.cells)]
            max_width = max([display_width(sample) for sample in samples] or [4])
            head = header[idx] if idx < len(header) else ""
            if "序号" in head:
                weight = 3.0
            elif all(is_short_value(sample) for sample in samples if sample):
                weight = 4.0
            else:
                weight = min(max(max_width / 2.2, 6.0), 16.0)
            weights.append(weight)
        total_weight = sum(weights) or 1
        widths = [max(650, int(total_width * weight / total_weight)) for weight in weights]

    diff = total_width - sum(widths)
    widths[-1] += diff
    return widths


def set_table_borders(table, outer_color: str = "8C8C8C", inner_color: str = "D9D9D9"):
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    for child_name in ("tblStyle", "tblLook"):
        child = tbl_pr.first_child_found_in(f"w:{child_name}")
        if child is not None:
            tbl_pr.remove(child)
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        is_outer = edge in {"top", "left", "bottom", "right"}
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), "6" if is_outer else "3")
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), outer_color if is_outer else inner_color)


def set_table_geometry(table, widths: list[int], total_width: int = 9070):
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total_width))
    tbl_w.set(qn("w:type"), "dxa")

    layout = tbl_pr.first_child_found_in("w:tblLayout")
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    tbl_grid = tbl.find(qn("w:tblGrid"))
    if tbl_grid is None:
        tbl_grid = OxmlElement("w:tblGrid")
        tbl.insert(1, tbl_grid)
    for child in list(tbl_grid):
        tbl_grid.remove(child)
    for width in widths:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        tbl_grid.append(grid_col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            width = widths[min(idx, len(widths) - 1)]
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")


def repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = tr_pr.first_child_found_in("w:tblHeader")
    if tbl_header is None:
        tbl_header = OxmlElement("w:tblHeader")
        tr_pr.append(tbl_header)
    tbl_header.set(qn("w:val"), "true")


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_end)
    set_run_font(run, size_pt=10.5)


def apply_page_geometry(section, preset: dict | None = None):
    page = (preset or BUILTIN_PRESETS["academic-paper"])["page"]
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width = Cm(page["width_cm"])
    section.page_height = Cm(page["height_cm"])
    section.top_margin = Cm(page["top_cm"])
    section.bottom_margin = Cm(page["bottom_cm"])
    section.left_margin = Cm(page["left_cm"])
    section.right_margin = Cm(page["right_cm"])
    section.header_distance = Cm(1.5)
    section.footer_distance = Cm(1.5)


def apply_reference_page_geometry(doc: Document, style_source: Path):
    try:
        source_doc = Document(style_source)
    except Exception:
        return
    if not source_doc.sections:
        return
    source = source_doc.sections[0]
    for section in doc.sections:
        section.orientation = source.orientation
        section.page_width = source.page_width
        section.page_height = source.page_height
        section.top_margin = source.top_margin
        section.bottom_margin = source.bottom_margin
        section.left_margin = source.left_margin
        section.right_margin = source.right_margin
        section.header_distance = source.header_distance
        section.footer_distance = source.footer_distance


def configure_styles(doc: Document, preset: dict):
    fonts = preset["fonts"]
    body = preset["body"]
    title_spec = preset["title"]
    headings = preset["headings"]
    normal = doc.styles["Normal"]
    set_style_font(normal, fonts["east_asia"], fonts["ascii"], body["size"])
    normal.paragraph_format.line_spacing = body["line"]
    normal.paragraph_format.first_line_indent = Cm(body["first_indent_cm"])
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(body["space_after"])

    paper_title = get_or_create_paragraph_style(doc, TITLE_STYLE_NAME)
    paper_title.base_style = normal
    title_style = get_or_create_paragraph_style(doc, "Title")
    for title in (title_style, paper_title):
        set_style_font(title, fonts["east_asia"], fonts["ascii"], title_spec["size"], title_spec["bold"])
        remove_paragraph_borders(title)
        title.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title.paragraph_format.first_line_indent = Cm(0)
        title.paragraph_format.space_before = Pt(0)
        title.paragraph_format.space_after = Pt(title_spec["space_after"])
        title.paragraph_format.line_spacing = title_spec["line"]

    heading_specs = [
        ("Heading 1", headings["h1"], True, 12, 6),
        ("Heading 2", headings["h2"], True, 9, 4),
        ("Heading 3", headings["h3"], False, 6, 3),
    ]
    for name, size, bold, before, after in heading_specs:
        style = get_or_create_paragraph_style(doc, name)
        set_style_font(style, fonts["east_asia"], fonts["ascii"], size, bold)
        remove_paragraph_borders(style)
        style.paragraph_format.first_line_indent = Cm(0)
        style.paragraph_format.left_indent = Cm(0)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.5

    for style_name in ["Caption", "Table Caption"]:
        if style_name in doc.styles:
            set_style_font(doc.styles[style_name], fonts["east_asia"], fonts["ascii"], 9)


def create_reference_docx(path: Path, title: str = DEFAULT_DOCUMENT_TITLE, preset: dict | None = None):
    preset = preset or BUILTIN_PRESETS["academic-paper"]
    doc = Document()
    for section in doc.sections:
        apply_page_geometry(section, preset)
        footer = section.footer
        footer.paragraphs[0].text = ""
        add_page_number(footer.paragraphs[0])

    configure_styles(doc, preset)

    doc.add_paragraph(title, style=TITLE_STYLE_NAME)
    doc.add_paragraph("摘要", style="Heading 1")
    doc.add_paragraph("正文样式示例。")
    doc.add_paragraph("1 数据来源与研究方法", style="Heading 1")
    doc.add_paragraph("1.1 数据来源与样本特征", style="Heading 2")
    doc.add_paragraph("1.1.1 示例三级标题", style="Heading 3")
    table = doc.add_table(rows=2, cols=3)
    table.style = "Table Grid"
    for row in table.rows:
        for cell in row.cells:
            cell.text = "示例"
            set_cell_margins(cell)
    doc.save(path)


def run_pandoc(
    source_md: Path,
    reference_docx: Path,
    out_docx: Path,
    pandoc: str,
    resource_paths: list[Path],
):
    resource_path = os.pathsep.join(str(path) for path in resource_paths)
    cmd = [
        pandoc,
        str(source_md),
        "-f",
        "markdown+simple_tables+pipe_tables+multiline_tables+grid_tables+east_asian_line_breaks",
        "-o",
        str(out_docx),
        "--reference-doc",
        str(reference_docx),
        "--resource-path",
        resource_path,
        "--wrap=none",
    ]
    subprocess.run(cmd, check=True, cwd=str(source_md.parent))


def convert_with_libreoffice_to_docx(source: Path, run_dir: Path, soffice: str) -> Path:
    input_copy = run_dir / f"source_input{source.suffix.lower()}"
    shutil.copy2(source, input_copy)
    out_dir = run_dir / "lo_docx"
    profile_dir = run_dir / f"lo_convert_profile_{os.getpid()}"
    out_dir.mkdir(parents=True, exist_ok=True)
    profile_dir.mkdir(parents=True, exist_ok=True)
    profile_url = "file:///" + str(profile_dir.resolve()).replace("\\", "/")
    subprocess.run(
        [
            soffice,
            f"-env:UserInstallation={profile_url}",
            "--headless",
            "--invisible",
            "--norestore",
            "--nodefault",
            "--nofirststartwizard",
            "--convert-to",
            "docx",
            "--outdir",
            str(out_dir),
            str(input_copy),
        ],
        check=True,
    )
    expected = out_dir / f"{input_copy.stem}.docx"
    if expected.exists():
        return expected
    outputs = sorted(out_dir.glob("*.docx"))
    if not outputs:
        raise RuntimeError(f"LibreOffice did not convert {source} to DOCX.")
    return outputs[0]


def copy_docx_style_parts(source_docx: Path, style_source: Path, out_docx: Path) -> Path:
    style_parts = {
        "word/styles.xml",
        "word/theme/theme1.xml",
        "word/fontTable.xml",
        "word/settings.xml",
    }
    with ZipFile(style_source, "r") as reference_zip:
        replacements = {
            name: reference_zip.read(name)
            for name in style_parts
            if name in reference_zip.namelist()
        }
    if not replacements:
        shutil.copy2(source_docx, out_docx)
        return out_docx
    with ZipFile(source_docx, "r") as zin, ZipFile(out_docx, "w") as zout:
        written = set()
        for item in zin.infolist():
            if item.filename in replacements:
                zout.writestr(item, replacements[item.filename])
                written.add(item.filename)
            else:
                zout.writestr(item, zin.read(item.filename))
        for name, data in replacements.items():
            if name not in written:
                zout.writestr(name, data)
    return out_docx


def is_dash_rule(line: str) -> bool:
    stripped = line.strip()
    return len(stripped) >= 5 and set(stripped) <= {"-"}


def is_simple_table_separator(line: str) -> bool:
    stripped = line.strip()
    return (
        len(stripped) >= 5
        and set(stripped) <= {"-", " "}
        and len(re.findall(r"-{3,}", line)) >= 2
    )


def separator_spans(line: str) -> list[tuple[int, int]]:
    return [(m.start(), m.end()) for m in re.finditer(r"-{3,}", line)]


def split_fixed_width_row(line: str, spans: list[tuple[int, int]]) -> list[str]:
    cells: list[str] = []
    for idx, (start, span_end) in enumerate(spans):
        end = spans[idx + 1][0] if idx + 1 < len(spans) else max(len(line), span_end)
        value = line[start:end] if start < len(line) else ""
        cells.append(re.sub(r"\s+", " ", value.strip()))
    return cells


def split_bold_header(line: str) -> list[str]:
    cells = re.findall(r"\*\*(.*?)\*\*", line)
    return [f"**{cell.strip()}**" for cell in cells if cell.strip()]


def plain_cell(cell: str) -> str:
    return re.sub(r"[*_`]", "", cell)


def split_numeric_prefix(value: str) -> tuple[str, str] | None:
    match = re.match(r"^(\d+(?:\.\d+)?)\s+(.+)$", value.strip())
    if not match:
        return None
    return match.group(1), match.group(2).strip()


def normalize_table_row(cells: list[str], header: list[str]) -> list[str]:
    cells = [re.sub(r"\s+", " ", cell.strip()) for cell in cells if cell.strip()]
    header_plain = [plain_cell(cell) for cell in header]

    if len(cells) < len(header) and header_plain and "序号" in header_plain[0] and cells:
        split = split_numeric_prefix(cells[0])
        if split:
            cells = [split[0], split[1], *cells[1:]]

    idx = 0
    while len(cells) < len(header) and idx < min(len(cells), len(header_plain)):
        name = header_plain[idx]
        if any(key in name for key in ("频次", "占比", "评分", "指数", "边权", "支持度")):
            split = split_numeric_prefix(cells[idx])
            if split:
                cells = [*cells[:idx], split[0], split[1], *cells[idx + 1 :]]
                idx += 2
                continue
        idx += 1

    if len(cells) > len(header):
        cells = cells[: len(header) - 1] + [" ".join(cells[len(header) - 1 :])]
    while len(cells) < len(header):
        cells.append("")
    return cells


def escape_pipe_cell(cell: str) -> str:
    return cell.replace("|", r"\|")


def split_table_row(line: str, col_count: int | None = None) -> list[str]:
    cells = [c.strip() for c in re.split(r"\s{2,}", line.strip()) if c.strip()]
    if col_count is not None:
        if len(cells) > col_count:
            cells = cells[: col_count - 1] + [" ".join(cells[col_count - 1 :])]
        while len(cells) < col_count:
            cells.append("")
    return cells


def table_to_pipe(block: list[str]) -> list[str]:
    content = [ln.rstrip() for ln in block if ln.strip() and not is_dash_rule(ln)]
    if len(content) < 2:
        return block

    separator_index = next((idx for idx, ln in enumerate(content) if is_simple_table_separator(ln)), None)
    if separator_index is not None and separator_index > 0:
        spans = separator_spans(content[separator_index])
        header = split_bold_header(content[separator_index - 1])
        if not header:
            header = split_fixed_width_row(content[separator_index - 1], spans)
        source_rows = content[separator_index + 1 :]
        rows = []
        for line in source_rows:
            if is_simple_table_separator(line):
                continue
            split_row = split_table_row(line)
            if len(split_row) < 2 and spans:
                split_row = split_fixed_width_row(line, spans)
            rows.append(normalize_table_row(split_row, header))
    else:
        header = split_table_row(content[0])
        source_rows = content[1:]
        rows = []
        for line in source_rows:
            if is_dash_rule(line) or set(line.strip()) <= {"-", " "}:
                continue
            rows.append(split_table_row(line, len(header)))

    header = [cell or f"列{idx + 1}" for idx, cell in enumerate(header)]
    if len(header) < 2:
        return block
    rows = [row[: len(header)] + [""] * max(0, len(header) - len(row)) for row in rows]
    if not rows:
        return block

    out = [
        "| " + " | ".join(escape_pipe_cell(cell) for cell in header) + " |",
        "| " + " | ".join(["---"] * len(header)) + " |",
    ]
    out.extend("| " + " | ".join(escape_pipe_cell(cell) for cell in row) + " |" for row in rows)
    return out


def preprocess_markdown(source: Path, output: Path):
    lines = read_text_compat(source).splitlines()
    result: list[str] = []
    i = 0
    while i < len(lines):
        if (
            i + 1 < len(lines)
            and (is_simple_table_separator(lines[i + 1]) or is_dash_rule(lines[i + 1]))
            and len(split_table_row(lines[i])) >= 2
        ):
            header = split_table_row(lines[i])
            rows: list[list[str]] = []
            j = i + 2
            while j < len(lines) and lines[j].strip():
                if is_simple_table_separator(lines[j]) or is_dash_rule(lines[j]):
                    j += 1
                    continue
                row = split_table_row(lines[j], len(header))
                if len([cell for cell in row if cell]) >= 2:
                    rows.append(normalize_table_row(row, header))
                    j += 1
                    continue
                break
            if rows:
                if result and result[-1].strip():
                    result.append("")
                result.extend(
                    [
                        "| " + " | ".join(escape_pipe_cell(cell) for cell in header) + " |",
                        "| " + " | ".join(["---"] * len(header)) + " |",
                    ]
                )
                result.extend("| " + " | ".join(escape_pipe_cell(cell) for cell in row) + " |" for row in rows)
                result.append("")
                i = j
                continue

        if is_dash_rule(lines[i]):
            j = i + 1
            block = [lines[i]]
            while j < len(lines):
                block.append(lines[j])
                if j > i and is_dash_rule(lines[j]):
                    break
                j += 1
            if len(block) >= 4 and is_dash_rule(block[-1]):
                converted = table_to_pipe(block)
                result.extend(converted)
                i = j + 1
                continue
        result.append(lines[i])
        i += 1

    output.write_text("\n".join(result) + "\n", encoding="utf-8")


def looks_like_caption(text: str) -> bool:
    stripped = text.strip()
    return bool(
        re.match(r"^(?:\u56fe|\u8868)\s*\d+(?:\s*[-\u2011\u2013\u2014\uff0d]\s*\d+)?", stripped)
        or re.match(r"^(?:figure|fig\.?|table)\s+\d+(?:[.\-:]\d+)?(?:$|\s*[:.\-]\s+|\s+)", stripped, re.IGNORECASE)
    )


def apply_body_style(doc: Document, para, preset: dict):
    body = preset["body"]
    para.style = doc.styles["Normal"]
    para.paragraph_format.first_line_indent = Cm(body["first_indent_cm"])
    para.paragraph_format.line_spacing = body["line"]
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.space_after = Pt(body["space_after"])


def force_normal_for_body_text_styles(doc: Document):
    """Convert leftover Word Body Text paragraph styles to Normal."""
    for p_style in doc.element.xpath(".//w:pStyle"):
        style_id = p_style.get(qn("w:val"))
        if style_id and style_id.startswith("BodyText"):
            p_style.set(qn("w:val"), "Normal")


def is_reference_heading(text: str) -> bool:
    stripped = text.strip()
    return stripped == "参考文献" or stripped.lower() in {"references", "bibliography", "works cited"}


def is_reference_entry(text: str) -> bool:
    return bool(re.match(r"^\[\d+\]", text.strip()))


def citation_superscript_size(preset: dict) -> float:
    body_size = float(preset["body"]["size"])
    return max(6.5, min(body_size * 0.75, body_size - 1.5))


def get_or_add_run_pr_element(r_element):
    r_pr = r_element.find(qn("w:rPr"))
    if r_pr is None:
        r_pr = OxmlElement("w:rPr")
        r_element.insert(0, r_pr)
    return r_pr


def set_run_element_superscript(r_element, size_pt: float | None = None):
    r_pr = get_or_add_run_pr_element(r_element)
    vert_align = r_pr.find(qn("w:vertAlign"))
    if vert_align is None:
        vert_align = OxmlElement("w:vertAlign")
        r_pr.append(vert_align)
    vert_align.set(qn("w:val"), "superscript")

    if size_pt is not None:
        half_points = str(int(round(size_pt * 2)))
        for tag in ("w:sz", "w:szCs"):
            size_node = r_pr.find(qn(tag))
            if size_node is None:
                size_node = OxmlElement(tag)
                r_pr.append(size_node)
            size_node.set(qn("w:val"), half_points)


def clone_text_run_element(base_r, text: str, superscript: bool = False, size_pt: float | None = None):
    new_r = deepcopy(base_r)
    for child in list(new_r):
        if child.tag != qn("w:rPr"):
            new_r.remove(child)
    if superscript:
        set_run_element_superscript(new_r, size_pt)
    text_node = OxmlElement("w:t")
    if text.startswith(" ") or text.endswith(" "):
        text_node.set(qn("xml:space"), "preserve")
    text_node.text = text
    new_r.append(text_node)
    return new_r


def superscript_inline_citations(paragraph, size_pt: float) -> int:
    changed = 0
    for run in list(paragraph.runs):
        text = run.text
        if not text or not CITATION_PATTERN.search(text):
            continue
        parent = run._r.getparent()
        insert_at = parent.index(run._r)
        cursor = 0
        for match in CITATION_PATTERN.finditer(text):
            if match.start() > cursor:
                parent.insert(insert_at, clone_text_run_element(run._r, text[cursor : match.start()]))
                insert_at += 1
            parent.insert(insert_at, clone_text_run_element(run._r, match.group(0), superscript=True, size_pt=size_pt))
            insert_at += 1
            changed += 1
            cursor = match.end()
        if cursor < len(text):
            parent.insert(insert_at, clone_text_run_element(run._r, text[cursor:]))
            insert_at += 1
        parent.remove(run._r)
    return changed


def post_process_docx(
    in_docx: Path,
    out_docx: Path,
    document_title: str,
    preset: dict,
    style_source: Path | None = None,
    preserve_reference_styles: bool = False,
) -> Path:
    working_docx = in_docx
    if style_source and preserve_reference_styles:
        working_docx = out_docx.with_name(f"{out_docx.stem}_style_source.docx")
        copy_docx_style_parts(in_docx, style_source, working_docx)
    doc = Document(working_docx)
    if style_source and preserve_reference_styles:
        apply_reference_page_geometry(doc, style_source)
    for section in doc.sections:
        if not preserve_reference_styles:
            apply_page_geometry(section, preset)
        footer = section.footer
        paragraph = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        paragraph.clear()
        add_page_number(paragraph)

    if not preserve_reference_styles:
        configure_styles(doc, preset)

    in_references = False
    body_size = preset["body"]["size"]
    title_size = preset["title"]["size"]
    heading_sizes = preset["headings"]
    citation_size = citation_superscript_size(preset)
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if not text:
            continue

        if i == 0 and text == document_title:
            try:
                para.style = doc.styles[TITLE_STYLE_NAME]
            except KeyError:
                para.style = doc.styles["Title"] if "Title" in doc.styles else doc.styles["Normal"]
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            para.paragraph_format.first_line_indent = Cm(0)
            remove_paragraph_borders(para)
        elif is_reference_heading(text):
            in_references = True
            para.style = doc.styles["Heading 1"]
        elif looks_like_caption(text):
            para.style = doc.styles["Normal"]
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            para.paragraph_format.first_line_indent = Cm(0)
            para.paragraph_format.space_before = Pt(3)
            para.paragraph_format.space_after = Pt(6)
            for run in para.runs:
                set_run_font(run, size_pt=9)
        elif in_references and re.match(r"^\[\d+\]", text):
            para.style = doc.styles["Normal"]
            para.paragraph_format.first_line_indent = Cm(-0.74)
            para.paragraph_format.left_indent = Cm(0.74)
            para.paragraph_format.line_spacing = 1.0
            para.paragraph_format.space_after = Pt(3)
        elif not para.style.name.startswith("Heading"):
            if preserve_reference_styles:
                para.style = doc.styles["Normal"]
            else:
                apply_body_style(doc, para, preset)

        if para.style.name.startswith("Heading"):
            para.paragraph_format.first_line_indent = Cm(0)
            para.paragraph_format.left_indent = Cm(0)
            remove_paragraph_borders(para)

        if not preserve_reference_styles:
            for run in para.runs:
                size = None
                bold = None
                if para.style.name in ("Title", TITLE_STYLE_NAME):
                    size, bold = title_size, True
                elif para.style.name == "Heading 1":
                    size, bold = heading_sizes["h1"], True
                elif para.style.name == "Heading 2":
                    size, bold = heading_sizes["h2"], True
                elif para.style.name == "Heading 3":
                    size, bold = heading_sizes["h3"], False
                elif in_references and re.match(r"^\[\d+\]", text):
                    size = min(body_size, 9.5)
                set_run_font(run, size_pt=size, bold=bold)

        if not in_references and not para.style.name.startswith("Heading") and not looks_like_caption(text):
            superscript_inline_citations(para, citation_size)

    for table in doc.tables:
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = False
        table.allow_autofit = False
        widths = preferred_column_widths(table)
        if widths:
            set_table_geometry(table, widths)
        set_table_borders(table)
        if table.rows:
            repeat_table_header(table.rows[0])
        col_count = len(table.columns)
        table_font_size = 6.4 if col_count >= 8 else 7.6 if col_count >= 5 else 9.0
        margin = 45 if col_count >= 8 else 65 if col_count >= 5 else 95
        for row_idx, row in enumerate(table.rows):
            for col_idx, cell in enumerate(row.cells):
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                set_cell_margins(cell, top=margin, start=margin, bottom=margin, end=margin)
                set_cell_shading(cell, "EDEDED" if row_idx == 0 else "FFFFFF")
                for para in cell.paragraphs:
                    remove_paragraph_style(para)
                    para.paragraph_format.first_line_indent = Cm(0)
                    para.paragraph_format.line_spacing = 1.0
                    para.paragraph_format.space_after = Pt(0)
                    para.paragraph_format.space_before = Pt(0)
                    text = para.text.strip()
                    if row_idx == 0 or col_idx == 0 and text.isdigit() or col_count >= 8 and col_idx > 0 or is_short_value(text):
                        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    else:
                        para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    for run in para.runs:
                        set_run_font(run, size_pt=table_font_size, bold=True if row_idx == 0 else None)

    force_normal_for_body_text_styles(doc)

    try:
        doc.save(out_docx)
        return out_docx
    except PermissionError:
        fallback = out_docx.with_name(f"{out_docx.stem}_new_{os.getpid()}{out_docx.suffix}")
        doc.save(fallback)
        return fallback


def resolve_tool(explicit: str | None, env_vars: list[str], names: list[str], candidates: list[Path]) -> str | None:
    values: list[str] = []
    if explicit:
        values.append(explicit)
    for env_var in env_vars:
        value = os.environ.get(env_var)
        if value:
            values.append(value)
    for value in values:
        path = Path(value)
        if path.exists():
            return str(path)
        found = shutil.which(value)
        if found:
            return found
    for name in names:
        found = shutil.which(name)
        if found:
            return found
    for candidate in candidates:
        if candidate.exists():
            return str(candidate)
    return None


def default_tool_candidates() -> dict[str, list[Path]]:
    home = Path.home()
    return {
        "pandoc": [
            Path(r"C:\Program Files\Pandoc\pandoc.exe"),
            home / r"AppData\Local\Pandoc\pandoc.exe",
        ],
        "libreoffice": [
            Path(r"D:\LibreOffice\program\soffice.com"),
            Path(r"D:\LibreOffice\program\soffice.exe"),
            Path(r"C:\Program Files\LibreOffice\program\soffice.com"),
            Path(r"C:\Program Files\LibreOffice\program\soffice.exe"),
            Path(r"C:\Program Files (x86)\LibreOffice\program\soffice.com"),
            Path(r"C:\Program Files (x86)\LibreOffice\program\soffice.exe"),
        ],
        "pdftoppm": [
            Path(r"D:\texlive\2026\bin\windows\pdftoppm.exe"),
            Path(r"D:\texlive\2025\bin\windows\pdftoppm.exe"),
            Path(r"C:\Program Files\poppler\Library\bin\pdftoppm.exe"),
            Path(r"C:\Program Files\poppler\bin\pdftoppm.exe"),
        ],
    }


def check_environment(args) -> dict:
    candidates = default_tool_candidates()
    report = {
        "pandoc": resolve_tool(args.pandoc, ["PANDOC"], ["pandoc"], candidates["pandoc"]),
        "libreoffice": resolve_tool(args.libreoffice, ["LIBREOFFICE", "SOFFICE"], ["soffice.com", "soffice"], candidates["libreoffice"]),
        "pdftoppm": resolve_tool(args.pdftoppm, ["PDFTOPPM"], ["pdftoppm"], candidates["pdftoppm"]),
    }
    report["missing"] = [name for name, value in report.items() if name in {"pandoc", "libreoffice", "pdftoppm"} and not value]
    report["install_hints"] = {
        "pandoc": "Install Pandoc from https://pandoc.org/installing.html or run: winget install JohnMacFarlane.Pandoc",
        "libreoffice": "Install LibreOffice from https://www.libreoffice.org/download/download-libreoffice/ or run: winget install TheDocumentFoundation.LibreOffice",
        "pdftoppm": "Install Poppler or TeX Live. With winget, try: winget install oschwartz10612.Poppler",
    }
    return report


def input_kind(source: Path) -> str:
    suffix = source.suffix.lower()
    if suffix in TEXT_INPUT_EXTENSIONS:
        return "text"
    if suffix in DOCX_INPUT_EXTENSIONS:
        return "docx"
    if suffix in LIBREOFFICE_INPUT_EXTENSIONS:
        return "libreoffice"
    raise ValueError(
        "Unsupported input format. Supported: "
        + ", ".join(sorted(TEXT_INPUT_EXTENSIONS | DOCX_INPUT_EXTENSIONS | LIBREOFFICE_INPUT_EXTENSIONS))
    )


def deep_merge(base: dict, override: dict) -> dict:
    merged = json.loads(json.dumps(base, ensure_ascii=False))
    for key, value in override.items():
        if isinstance(value, dict) and isinstance(merged.get(key), dict):
            merged[key] = deep_merge(merged[key], value)
        else:
            merged[key] = value
    return merged


def load_custom_presets() -> dict[str, dict]:
    presets: dict[str, dict] = {}
    if not CUSTOM_PRESETS_DIR.exists():
        return presets
    for path in CUSTOM_PRESETS_DIR.glob("*.json"):
        try:
            presets[path.stem] = json.loads(path.read_text(encoding="utf-8-sig"))
        except Exception:
            continue
    return presets


def contains_cjk(text: str) -> bool:
    return any("\u4e00" <= ch <= "\u9fff" for ch in text)


def looks_english(text: str) -> bool:
    letters = sum(1 for ch in text if ch.isascii() and ch.isalpha())
    cjk = sum(1 for ch in text if "\u4e00" <= ch <= "\u9fff")
    return letters >= 8 and letters > cjk * 2


def is_english_preset_name(name: str | None) -> bool:
    return bool(name and name.startswith("english-"))


def infer_preset_name(source: Path, title: str) -> str:
    text = f"{source.stem} {title}".lower()
    is_english = looks_english(text) and not contains_cjk(text)
    if any(key in text for key in ["论文", "研究", "paper", "thesis"]):
        return "english-academic" if is_english else "academic-paper"
    if any(key in text for key in ["汇报", "briefing", "简报"]):
        return "english-briefing" if is_english else "briefing"
    if any(key in text for key in ["演讲", "讲话", "发言", "speech"]):
        return "english-speech" if is_english else "speech"
    if any(key in text for key in ["报告", "report"]):
        return "english-report" if is_english else "report"
    return "english-general" if is_english else "general"


def resolve_preset(args, source: Path, title: str, style_source: Path | None) -> tuple[str, dict, bool]:
    custom_presets = load_custom_presets()
    available = {**BUILTIN_PRESETS, **custom_presets}
    preset_name = args.preset
    preserve_reference_styles = False
    if preset_name == "auto":
        preset_name = "reference" if style_source else infer_preset_name(source, title)
    if preset_name == "reference":
        preserve_reference_styles = bool(style_source)
    if preset_name not in available:
        raise ValueError(f"Unknown preset '{preset_name}'. Available presets: {', '.join(sorted(available))}")
    preset = available[preset_name]
    if args.preset_file:
        user_preset = json.loads(args.preset_file.read_text(encoding="utf-8-sig"))
        preset = deep_merge(preset, user_preset)
    if args.save_preset_name:
        CUSTOM_PRESETS_DIR.mkdir(parents=True, exist_ok=True)
        preset_path = CUSTOM_PRESETS_DIR / f"{args.save_preset_name}.json"
        preset_path.write_text(json.dumps(preset, ensure_ascii=False, indent=2), encoding="utf-8")
        preset_name = args.save_preset_name
    return preset_name, preset, preserve_reference_styles


def require_environment(report: dict, render: bool, kind: str):
    missing = []
    if kind == "text" and not report["pandoc"]:
        missing.append("pandoc")
    if (render or kind == "libreoffice") and not report["libreoffice"]:
        missing.append("libreoffice")
    if render and not report["pdftoppm"]:
        missing.append("pdftoppm")
    if not missing:
        return
    lines = ["Missing required document tooling:"]
    for name in missing:
        lines.append(f"- {name}: {report['install_hints'][name]}")
    raise RuntimeError("\n".join(lines))


def infer_title(source: Path, explicit_title: str | None) -> str:
    if explicit_title:
        return explicit_title.strip()
    try:
        if source.suffix.lower() == ".docx":
            doc = Document(source)
            for para in doc.paragraphs[:20]:
                text = para.text.strip()
                if text:
                    return text
        elif source.suffix.lower() in TEXT_INPUT_EXTENSIONS:
            for line in read_text_compat(source).splitlines():
                text = line.strip()
                if not text:
                    continue
                text = re.sub(r"^#+\s*", "", text).strip()
                text = text.strip("*_`# ")
                if text:
                    return text
    except OSError:
        pass
    return source.stem


def final_docx_path(source: Path, args) -> Path:
    if args.final:
        return args.final
    output_dir = args.output_dir or source.parent / DEFAULT_OUTPUT_DIR_NAME
    resolved_preset = getattr(args, "resolved_preset", args.preset)
    suffix = "_formatted" if is_english_preset_name(resolved_preset) else "_格式化"
    default_name = f"{source.stem}.docx" if source.stem.endswith(("_格式化", "_formatted")) else f"{source.stem}{suffix}.docx"
    output_name = args.output_name or default_name
    if not output_name.lower().endswith(".docx"):
        output_name += ".docx"
    return output_dir / output_name


def render_docx(docx_path: Path, run_dir: Path, soffice: str, pdftoppm: str) -> dict:
    pdf_dir = run_dir / "render_pdf"
    png_dir = run_dir / "render_png"
    profile_dir = run_dir / f"lo_profile_{os.getpid()}"
    pdf_dir.mkdir(parents=True, exist_ok=True)
    png_dir.mkdir(parents=True, exist_ok=True)
    profile_dir.mkdir(parents=True, exist_ok=True)
    for stale_pdf in pdf_dir.glob("*.pdf"):
        stale_pdf.unlink(missing_ok=True)
    for stale_png in png_dir.glob("page-*.png"):
        stale_png.unlink(missing_ok=True)
    profile_url = "file:///" + str(profile_dir.resolve()).replace("\\", "/")

    subprocess.run(
        [
            soffice,
            f"-env:UserInstallation={profile_url}",
            "--headless",
            "--invisible",
            "--norestore",
            "--nodefault",
            "--nofirststartwizard",
            "--convert-to",
            "pdf",
            "--outdir",
            str(pdf_dir),
            str(docx_path),
        ],
        check=True,
    )
    pdf_path = pdf_dir / f"{docx_path.stem}.pdf"
    if not pdf_path.exists():
        pdfs = sorted(pdf_dir.glob("*.pdf"))
        if not pdfs:
            raise RuntimeError("LibreOffice did not produce a PDF render.")
        pdf_path = pdfs[0]

    subprocess.run([pdftoppm, "-r", "150", "-png", str(pdf_path), str(png_dir / "page")], check=True)
    pages = sorted(png_dir.glob("page-*.png"))
    if not pages:
        raise RuntimeError("pdftoppm did not produce PNG page renders.")
    return {"pdf": str(pdf_path), "png_dir": str(png_dir), "pages": len(pages)}


def audit_body_text_styles(docx_path: Path) -> dict:
    counts: dict[str, int] = {}
    with ZipFile(docx_path) as archive:
        for name in archive.namelist():
            if not (name.startswith("word/") and name.endswith(".xml")) or name == "word/styles.xml":
                continue
            xml = archive.read(name).decode("utf-8", errors="ignore")
            for style_id in re.findall(r'<w:pStyle[^>]+w:val="([^"]+)"', xml):
                if style_id.startswith("BodyText"):
                    counts[style_id] = counts.get(style_id, 0) + 1
    return {"body_text_styles": counts, "ok": not counts}


def audit_inline_citation_superscripts(docx_path: Path) -> dict:
    import xml.etree.ElementTree as ET

    ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
    result = {
        "inline_citation_count": 0,
        "missing_superscript_count": 0,
        "examples": [],
        "ok": True,
    }
    with ZipFile(docx_path) as archive:
        xml = archive.read("word/document.xml")
    root = ET.fromstring(xml)
    in_references = False

    for para in root.findall(".//w:p", ns):
        runs = []
        para_text = ""
        for run in para.findall("./w:r", ns):
            text = "".join(node.text or "" for node in run.findall(".//w:t", ns))
            if not text:
                continue
            para_text += text
            runs.append((run, text))

        stripped = para_text.strip()
        if not stripped:
            continue
        if is_reference_heading(stripped):
            in_references = True
            continue
        if in_references or is_reference_entry(stripped):
            continue

        for run, text in runs:
            for match in CITATION_PATTERN.finditer(text):
                result["inline_citation_count"] += 1
                r_pr = run.find("./w:rPr", ns)
                vert_align = r_pr.find("./w:vertAlign", ns) if r_pr is not None else None
                value = vert_align.get(qn("w:val")) if vert_align is not None else None
                if value not in {"superscript", "super"}:
                    result["missing_superscript_count"] += 1
                    if len(result["examples"]) < 10:
                        result["examples"].append({"citation": match.group(0), "paragraph": stripped[:140]})

    result["ok"] = result["missing_superscript_count"] == 0
    return result


def add_issue(items: list[dict], code: str, message: str, **extra):
    issue = {"code": code, "message": message}
    issue.update(extra)
    items.append(issue)


def is_under_path(path: Path, parent: Path | None) -> bool:
    if parent is None:
        return False
    try:
        path.resolve().relative_to(parent.resolve())
        return True
    except ValueError:
        return False


def detect_build_root(path: Path) -> Path | None:
    parts = path.resolve().parts
    if DEFAULT_BUILD_DIR_NAME not in parts:
        return None
    index = parts.index(DEFAULT_BUILD_DIR_NAME)
    return Path(*parts[: index + 1])


def png_non_white_ratio(path: Path) -> float | None:
    try:
        from PIL import Image
    except Exception:
        return None
    try:
        with Image.open(path) as image:
            rgb = image.convert("RGB")
            width, height = rgb.size
            if width <= 0 or height <= 0:
                return None
            step = max(1, int(((width * height) / 220000) ** 0.5))
            total = 0
            ink = 0
            for y in range(0, height, step):
                for x in range(0, width, step):
                    r, g, b = rgb.getpixel((x, y))
                    if r < 245 or g < 245 or b < 245:
                        ink += 1
                    total += 1
            return ink / total if total else None
    except Exception:
        return None


def audit_image_relationships(archive: ZipFile) -> tuple[list[str], list[str]]:
    names = set(archive.namelist())
    missing: list[str] = []
    external: list[str] = []
    rel_files = [name for name in names if name.startswith("word/_rels/") and name.endswith(".rels")]
    for rel_file in rel_files:
        try:
            xml = archive.read(rel_file).decode("utf-8", errors="ignore")
        except KeyError:
            continue
        source_part = rel_file[len("word/_rels/") : -len(".rels")]
        source_dir = posixpath.dirname(posixpath.join("word", source_part))
        for rel_match in re.finditer(r"<Relationship\b[^>]+>", xml):
            rel_xml = rel_match.group(0)
            if "/image" not in rel_xml:
                continue
            target_match = re.search(r'Target="([^"]+)"', rel_xml)
            if not target_match:
                continue
            target = target_match.group(1).replace("\\", "/")
            if re.search(r'TargetMode="External"', rel_xml) or re.match(r"^[a-zA-Z][a-zA-Z0-9+.-]*:", target):
                external.append(target)
                continue
            target_path = posixpath.normpath(posixpath.join(source_dir, target.lstrip("/")))
            if target_path not in names:
                missing.append(target_path)
    return sorted(set(missing)), sorted(set(external))


def audit_docx_quality(
    docx_path: Path,
    build_root: Path | None = None,
    render_report: dict | None = None,
    render_expected: bool = False,
) -> dict:
    findings: list[dict] = []
    warnings: list[dict] = []
    metrics: dict = {
        "path": str(docx_path),
        "exists": docx_path.exists(),
        "size_bytes": docx_path.stat().st_size if docx_path.exists() else 0,
    }

    if not docx_path.exists():
        add_issue(findings, "missing-final", "Final DOCX does not exist.")
        return {"ok": False, "findings": findings, "warnings": warnings, "metrics": metrics}
    if docx_path.suffix.lower() != ".docx":
        add_issue(findings, "not-docx", "Final deliverable must be a .docx file.")
    if metrics["size_bytes"] < 6000:
        add_issue(findings, "too-small", "DOCX is unexpectedly small and may be empty or corrupt.", size_bytes=metrics["size_bytes"])
    if is_under_path(docx_path, build_root):
        add_issue(findings, "final-in-build-dir", "Final deliverable is inside the intermediate build directory.", build_root=str(build_root))

    try:
        with ZipFile(docx_path) as archive:
            bad_member = archive.testzip()
            if bad_member:
                add_issue(findings, "bad-docx-member", "DOCX package contains a corrupt internal member.", member=bad_member)
            missing_images, external_images = audit_image_relationships(archive)
            metrics["missing_image_count"] = len(missing_images)
            metrics["external_image_count"] = len(external_images)
            if missing_images:
                add_issue(findings, "missing-image-parts", "DOCX contains image relationships whose media files are missing.", images=missing_images[:10])
            if external_images:
                add_issue(warnings, "external-images", "DOCX references external images; confirm they render in the target environment.", images=external_images[:10])
    except BadZipFile:
        add_issue(findings, "bad-zip", "DOCX is not a valid ZIP/OOXML package.")
    except Exception as exc:
        add_issue(findings, "zip-audit-error", f"Could not inspect DOCX package: {exc}")

    try:
        doc = Document(docx_path)
        text_chars = sum(len(paragraph.text.strip()) for paragraph in doc.paragraphs)
        table_count = len(doc.tables)
        metrics.update(
            {
                "paragraph_count": len(doc.paragraphs),
                "table_count": table_count,
                "section_count": len(doc.sections),
                "text_chars": text_chars,
            }
        )
        if text_chars == 0 and table_count == 0:
            add_issue(findings, "empty-document", "DOCX has no paragraph text or tables.")
        if len(doc.sections) == 0:
            add_issue(findings, "no-sections", "DOCX has no sections.")
        malformed_tables = []
        wide_tables = []
        for index, table in enumerate(doc.tables, start=1):
            row_count = len(table.rows)
            col_count = len(table.columns)
            if row_count == 0 or col_count == 0:
                malformed_tables.append({"table": index, "rows": row_count, "columns": col_count})
            if col_count >= 8:
                wide_tables.append({"table": index, "columns": col_count})
        metrics["malformed_table_count"] = len(malformed_tables)
        metrics["wide_table_count"] = len(wide_tables)
        if malformed_tables:
            add_issue(findings, "malformed-tables", "Some tables have zero rows or zero columns.", tables=malformed_tables[:10])
        if wide_tables:
            add_issue(warnings, "wide-tables", "Wide tables should be visually checked after rendering.", tables=wide_tables[:10])
    except Exception as exc:
        add_issue(findings, "docx-open-error", f"Could not open DOCX with python-docx: {exc}")

    try:
        style_audit = audit_body_text_styles(docx_path)
        metrics["body_text_styles"] = style_audit["body_text_styles"]
        if not style_audit["ok"]:
            add_issue(findings, "body-text-style", "DOCX still contains Word BodyText paragraph styles.", styles=style_audit["body_text_styles"])
    except Exception as exc:
        add_issue(findings, "style-audit-error", f"Could not inspect paragraph style drift: {exc}")

    try:
        citation_audit = audit_inline_citation_superscripts(docx_path)
        metrics["inline_citation_count"] = citation_audit["inline_citation_count"]
        metrics["inline_citation_missing_superscript_count"] = citation_audit["missing_superscript_count"]
        if not citation_audit["ok"]:
            add_issue(
                findings,
                "inline-citation-superscript",
                "Inline numeric citations in the document body must be superscripted.",
                examples=citation_audit["examples"],
            )
    except Exception as exc:
        add_issue(findings, "citation-audit-error", f"Could not inspect inline citation superscripts: {exc}")

    if render_expected and not render_report:
        add_issue(findings, "render-missing", "Render QA was required but no render report is available.")
    if render_report:
        metrics["render_pages"] = render_report.get("pages", 0)
        pdf_path = Path(render_report.get("pdf", ""))
        png_dir = Path(render_report.get("png_dir", ""))
        if not pdf_path.exists() or pdf_path.stat().st_size == 0:
            add_issue(findings, "render-pdf-missing", "LibreOffice render PDF is missing or empty.", pdf=str(pdf_path))
        pages = sorted(png_dir.glob("page-*.png")) if png_dir.exists() else []
        if not pages:
            add_issue(findings, "render-png-missing", "No PNG page renders were produced.", png_dir=str(png_dir))
        elif render_report.get("pages") and len(pages) != render_report.get("pages"):
            add_issue(warnings, "render-page-count-mismatch", "Rendered PNG count differs from render report.", report_pages=render_report.get("pages"), png_pages=len(pages))
        density_values = []
        unknown_density = 0
        for page_index, page in enumerate(pages, start=1):
            if page.stat().st_size < 1024:
                add_issue(findings, "render-page-too-small", "Rendered page PNG is unexpectedly small.", page=page_index, path=str(page), size_bytes=page.stat().st_size)
                continue
            ratio = png_non_white_ratio(page)
            if ratio is None:
                unknown_density += 1
                continue
            density_values.append(round(ratio, 6))
            if ratio < 0.0002:
                add_issue(findings, "blank-render-page", "Rendered page appears blank or nearly blank.", page=page_index, non_white_ratio=round(ratio, 6))
        metrics["render_non_white_ratios"] = density_values
        if pages and unknown_density == len(pages):
            add_issue(warnings, "render-density-unchecked", "Could not estimate PNG text density; install Pillow for stronger blank-page checks.")
    elif not render_expected:
        add_issue(warnings, "render-skipped", "Render QA was skipped; visually inspect the DOCX before final delivery if quality matters.")

    return {"ok": not findings, "findings": findings, "warnings": warnings, "metrics": metrics}


def repair_docx_quality(docx_path: Path, preset: dict, audit: dict) -> list[str]:
    codes = {item.get("code") for item in audit.get("findings", [])}
    repairable = codes & {"body-text-style", "malformed-tables", "inline-citation-superscript"}
    if not repairable:
        return []
    doc = Document(docx_path)
    actions: list[str] = []
    if "body-text-style" in repairable:
        force_normal_for_body_text_styles(doc)
        actions.append("converted BodyText paragraph styles to Normal")
    if "malformed-tables" in repairable:
        for table in doc.tables:
            if table.rows and table.columns:
                table.alignment = WD_TABLE_ALIGNMENT.CENTER
                table.autofit = False
                table.allow_autofit = False
                widths = preferred_column_widths(table)
                if widths:
                    set_table_geometry(table, widths)
                set_table_borders(table)
        actions.append("reapplied table geometry and borders")
    if "inline-citation-superscript" in repairable:
        in_references = False
        citation_size = citation_superscript_size(preset)
        count = 0
        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if not text:
                continue
            if is_reference_heading(text):
                in_references = True
                continue
            if in_references or is_reference_entry(text):
                continue
            count += superscript_inline_citations(paragraph, citation_size)
        if count:
            actions.append(f"converted {count} inline citation markers to superscript")
    if actions:
        doc.save(docx_path)
    return actions


def run_quality_review_loop(
    final_path: Path,
    build_root: Path | None,
    run_dir: Path,
    preset: dict,
    qa_iterations: int,
    render: bool,
    env_report: dict,
) -> tuple[dict, dict | None, list[dict]]:
    iterations: list[dict] = []
    render_report = None
    attempts = max(0, qa_iterations) + 1
    for attempt in range(attempts):
        if render:
            render_report = render_docx(final_path, run_dir, env_report["libreoffice"], env_report["pdftoppm"])
        audit = audit_docx_quality(final_path, build_root=build_root, render_report=render_report, render_expected=render)
        if audit["ok"] or attempt >= qa_iterations:
            return audit, render_report, iterations
        actions = repair_docx_quality(final_path, preset, audit)
        iterations.append({"iteration": attempt + 1, "actions": actions, "remaining_findings": audit["findings"]})
        if not actions:
            return audit, render_report, iterations
    return audit, render_report, iterations


def main():
    parser = argparse.ArgumentParser(description="Format or reformat a Chinese or English DOCX deliverable.")
    parser.add_argument("--source", type=Path, help="Source manuscript: .md, .markdown, .txt, .docx, .doc, .rtf, or .odt.")
    parser.add_argument("--title", help="Document title. Defaults to the first non-empty text line or Word paragraph.")
    parser.add_argument("--reference", type=Path, default=ASSET_REFERENCE_DOCX, help="Pandoc reference.docx.")
    parser.add_argument("--style-source", type=Path, help="Reference DOCX whose styles/page layout should be used as the formatting model.")
    parser.add_argument("--preset", default="auto", help="Format preset: auto, general, academic-paper, report, briefing, speech, english-general, english-academic, english-report, english-briefing, english-speech, reference, or a saved custom preset.")
    parser.add_argument("--preset-file", type=Path, help="JSON preset override file.")
    parser.add_argument("--save-preset-name", help="Save the resolved preset under this custom name for later reuse.")
    parser.add_argument("--output-dir", type=Path, help="Directory for final DOCX only. Defaults to <source>/docx_output.")
    parser.add_argument("--output-name", help="Final DOCX file name. Defaults to <source-stem>_格式化.docx for Chinese presets or <source-stem>_formatted.docx for English presets.")
    parser.add_argument("--final", type=Path, help="Exact final DOCX path; overrides --output-dir and --output-name.")
    parser.add_argument("--build-dir", type=Path, help="Directory for intermediate files. Defaults to <source>/_docx_style_build.")
    parser.add_argument("--resource-path", type=Path, action="append", default=[], help="Extra Pandoc resource path for images; used for text/Markdown inputs.")
    parser.add_argument("--pandoc", help="Explicit pandoc executable path.")
    parser.add_argument("--libreoffice", help="Explicit soffice/soffice.com executable path.")
    parser.add_argument("--pdftoppm", help="Explicit pdftoppm executable path.")
    parser.add_argument("--check-env", action="store_true", help="Only check local tooling and print JSON.")
    parser.add_argument("--skip-render", action="store_true", help="Skip LibreOffice/PDF/PNG render QA.")
    parser.add_argument("--no-overwrite", action="store_true", help="Fail if the final DOCX already exists.")
    parser.add_argument("--audit-only", action="store_true", help="Audit an existing DOCX without reformatting it.")
    parser.add_argument("--strict", action="store_true", help="Return a non-zero exit code if quality_audit.ok is false.")
    parser.add_argument("--qa-iterations", type=int, default=2, help="Maximum automatic repair iterations before final quality audit.")
    args = parser.parse_args()

    env_report = check_environment(args)
    if args.check_env:
        if args.source and args.source.exists():
            kind = input_kind(args.source)
            render = not args.skip_render
            required = []
            if kind == "text":
                required.append("pandoc")
            if render or kind == "libreoffice":
                required.append("libreoffice")
            if render:
                required.append("pdftoppm")
            env_report["input_kind"] = kind
            env_report["required"] = required
            env_report["required_missing"] = [name for name in required if not env_report.get(name)]
        print(json.dumps(env_report, ensure_ascii=False, indent=2))
        return 0 if not env_report.get("required_missing", env_report["missing"]) else 1

    if not args.source:
        parser.error("--source is required unless --check-env is used.")
    source = args.source.resolve()
    if not source.exists():
        raise FileNotFoundError(source)

    render = not args.skip_render

    if args.audit_only:
        if source.suffix.lower() != ".docx":
            raise ValueError("--audit-only requires a .docx source.")
        require_environment(env_report, render, "docx")
        build_root = args.build_dir.resolve() if args.build_dir else detect_build_root(source)
        run_root = args.build_dir or source.parent / DEFAULT_BUILD_DIR_NAME
        run_id = "audit_" + datetime.now().strftime("%Y%m%d_%H%M%S") + f"_{os.getpid()}"
        run_dir = run_root / run_id
        run_dir.mkdir(parents=True, exist_ok=True)
        render_report = render_docx(source, run_dir, env_report["libreoffice"], env_report["pdftoppm"]) if render else None
        quality_audit = audit_docx_quality(source, build_root=build_root, render_report=render_report, render_expected=render)
        report = {
            "final_docx": str(source),
            "build_dir": str(run_dir),
            "input_kind": "docx",
            "environment": env_report,
            "style_audit": {
                "body_text_styles": quality_audit.get("metrics", {}).get("body_text_styles", {}),
                "ok": not quality_audit.get("metrics", {}).get("body_text_styles", {}),
            },
            "quality_audit": quality_audit,
            "render": render_report,
        }
        (run_dir / "run_report.json").write_text(json.dumps(report, ensure_ascii=False, indent=2), encoding="utf-8")
        print(json.dumps(report, ensure_ascii=False, indent=2))
        return 2 if args.strict and not quality_audit["ok"] else 0

    kind = input_kind(source)
    require_environment(env_report, render, kind)

    document_title = infer_title(source, args.title)
    style_source = args.style_source.resolve() if args.style_source else None
    if style_source and not style_source.exists():
        raise FileNotFoundError(style_source)
    preset_name, preset, preserve_reference_styles = resolve_preset(args, source, document_title, style_source)
    args.resolved_preset = preset_name
    build_root = args.build_dir or source.parent / DEFAULT_BUILD_DIR_NAME
    run_id = datetime.now().strftime("%Y%m%d_%H%M%S") + f"_{os.getpid()}"
    run_dir = build_root / run_id
    run_dir.mkdir(parents=True, exist_ok=True)

    final_path = final_docx_path(source, args).resolve()
    final_path.parent.mkdir(parents=True, exist_ok=True)
    if final_path.exists() and args.no_overwrite:
        raise FileExistsError(final_path)

    reference_docx = style_source if style_source and kind == "text" else args.reference
    if not reference_docx.exists():
        reference_docx = run_dir / "reference.docx"
        create_reference_docx(reference_docx, document_title, preset)

    staged_docx = run_dir / "formatted.docx"

    if kind == "text":
        processed_md = run_dir / "processed_paper.md"
        intermediate_docx = run_dir / "pandoc.docx"
        resource_paths = [source.parent, *args.resource_path]
        preprocess_markdown(source, processed_md)
        run_pandoc(processed_md, reference_docx, intermediate_docx, env_report["pandoc"], resource_paths)
        source_docx = intermediate_docx
    elif kind == "docx":
        source_docx = run_dir / "source.docx"
        shutil.copy2(source, source_docx)
    else:
        source_docx = convert_with_libreoffice_to_docx(source, run_dir, env_report["libreoffice"])

    generated_docx = post_process_docx(
        source_docx,
        staged_docx,
        document_title,
        preset,
        style_source=style_source,
        preserve_reference_styles=preserve_reference_styles,
    )
    shutil.copy2(generated_docx, final_path)

    quality_audit, render_report, qa_iterations = run_quality_review_loop(
        final_path,
        build_root=build_root,
        run_dir=run_dir,
        preset=preset,
        qa_iterations=args.qa_iterations,
        render=render,
        env_report=env_report,
    )
    audit = {
        "body_text_styles": quality_audit.get("metrics", {}).get("body_text_styles", {}),
        "ok": not quality_audit.get("metrics", {}).get("body_text_styles", {}),
    }

    report = {
        "final_docx": str(final_path),
        "build_dir": str(run_dir),
        "reference_docx": str(reference_docx),
        "input_kind": kind,
        "preset": preset_name,
        "style_source": str(style_source) if style_source else None,
        "preserve_reference_styles": preserve_reference_styles,
        "source_docx": str(source_docx),
        "environment": env_report,
        "style_audit": audit,
        "quality_audit": quality_audit,
        "qa_iterations": qa_iterations,
        "render": render_report,
    }
    (run_dir / "run_report.json").write_text(json.dumps(report, ensure_ascii=False, indent=2), encoding="utf-8")
    print(json.dumps(report, ensure_ascii=False, indent=2))
    return 2 if args.strict and not quality_audit["ok"] else 0


if __name__ == "__main__":
    sys.exit(main())
