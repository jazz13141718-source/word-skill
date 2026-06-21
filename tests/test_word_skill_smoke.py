# -*- coding: utf-8 -*-
from __future__ import annotations

import importlib.util
import os
import py_compile
import shutil
import tempfile
from pathlib import Path
from types import SimpleNamespace
from zipfile import ZipFile
from xml.etree import ElementTree as ET

from docx import Document


REPO_ROOT = Path(__file__).resolve().parents[1]
SKILL_DIR = REPO_ROOT / "word-skill"
FORMATTER = SKILL_DIR / "scripts" / "format_docx_document.py"


def require(condition: bool, message: str):
    if not condition:
        raise AssertionError(message)


def validate_skill_metadata():
    skill_md = SKILL_DIR / "SKILL.md"
    text = skill_md.read_text(encoding="utf-8-sig")
    require(text.startswith("---"), "SKILL.md must start with YAML frontmatter.")
    require("\nname: word-skill\n" in text, "SKILL.md must declare name: word-skill.")
    require("\ndescription:" in text, "SKILL.md must include a description.")
    for relative in [
        "agents/generic.md",
        "agents/openai.yaml",
        "assets/reference.docx",
        "references/environment-and-output.md",
        "scripts/format_docx_document.py",
        "scripts/install_word_skill.py",
        "scripts/uninstall_word_skill.py",
    ]:
        require((SKILL_DIR / relative).exists(), f"Missing required file: {relative}")
    require((REPO_ROOT / "AGENTS.md").exists(), "Missing repository-level AGENTS.md.")


def validate_agent_portability_text():
    forbidden = [
        "C:\\Users\\ADMIN",
        "codex-runtimes",
        "codex-primary-runtime",
    ]
    text_files = [
        REPO_ROOT / "README.md",
        REPO_ROOT / "AGENTS.md",
        SKILL_DIR / "SKILL.md",
        SKILL_DIR / "agents" / "generic.md",
        SKILL_DIR / "references" / "environment-and-output.md",
        SKILL_DIR / "scripts" / "format_docx_document.py",
        SKILL_DIR / "scripts" / "install_word_skill.py",
        SKILL_DIR / "scripts" / "uninstall_word_skill.py",
    ]
    for path in text_files:
        text = path.read_text(encoding="utf-8-sig", errors="ignore")
        for needle in forbidden:
            require(needle not in text, f"Found non-portable local path marker {needle!r} in {path}")


def compile_scripts():
    for script in (SKILL_DIR / "scripts").glob("*.py"):
        py_compile.compile(str(script), doraise=True)


def load_formatter():
    spec = importlib.util.spec_from_file_location("format_docx_document", FORMATTER)
    module = importlib.util.module_from_spec(spec)
    assert spec.loader is not None
    spec.loader.exec_module(module)
    return module


def load_script_module(path: Path, module_name: str):
    spec = importlib.util.spec_from_file_location(module_name, path)
    module = importlib.util.module_from_spec(spec)
    assert spec.loader is not None
    spec.loader.exec_module(module)
    return module


def install_target_resolution_smoke():
    installer = load_script_module(SKILL_DIR / "scripts" / "install_word_skill.py", "install_word_skill")
    uninstaller = load_script_module(SKILL_DIR / "scripts" / "uninstall_word_skill.py", "uninstall_word_skill")
    tmp_root = Path(tempfile.mkdtemp(prefix="word_skill_target_"))
    old_env = {name: os.environ.get(name) for name in ["WORD_SKILL_HOME", "AGENT_SKILLS_DIR", "CODEX_HOME"]}
    try:
        for name in old_env:
            os.environ.pop(name, None)

        os.environ["AGENT_SKILLS_DIR"] = str(tmp_root / "agent-skills")
        expected = tmp_root / "agent-skills" / "word-skill"
        require(installer.resolve_target(None) == expected, "Installer did not honor AGENT_SKILLS_DIR.")
        require(uninstaller.resolve_target(None) == expected, "Uninstaller did not honor AGENT_SKILLS_DIR.")

        os.environ["WORD_SKILL_HOME"] = str(tmp_root / "exact" / "word-skill")
        exact = tmp_root / "exact" / "word-skill"
        require(installer.resolve_target(None) == exact, "Installer did not honor exact WORD_SKILL_HOME.")
        require(uninstaller.resolve_target(None) == exact, "Uninstaller did not honor exact WORD_SKILL_HOME.")

        os.environ["WORD_SKILL_HOME"] = str(tmp_root / "word-skill-parent")
        expected_home_parent = tmp_root / "word-skill-parent" / "word-skill"
        require(installer.resolve_target(None) == expected_home_parent, "Installer did not normalize parent WORD_SKILL_HOME.")
        require(uninstaller.resolve_target(None) == expected_home_parent, "Uninstaller did not normalize parent WORD_SKILL_HOME.")

        explicit_parent = tmp_root / "explicit-skills"
        require(installer.resolve_target(explicit_parent) == explicit_parent / "word-skill", "Installer explicit parent target failed.")
        require(uninstaller.resolve_target(explicit_parent) == explicit_parent / "word-skill", "Uninstaller explicit parent target failed.")
    finally:
        for name, value in old_env.items():
            if value is None:
                os.environ.pop(name, None)
            else:
                os.environ[name] = value
        shutil.rmtree(tmp_root, ignore_errors=True)


def citation_superscript_smoke():
    formatter = load_formatter()
    tmp_root = Path(tempfile.mkdtemp(prefix="word_skill_smoke_"))
    try:
        source = tmp_root / "source.docx"
        output = tmp_root / "output.docx"
        doc = Document()
        doc.add_paragraph("Test Title")
        doc.add_paragraph("Body citation[1-2] and another[3\uff0c4]. Interval[0,1] should stay normal.")
        doc.add_paragraph("\u53c2\u8003\u6587\u732e")
        doc.add_paragraph("[1] Reference entry should not be superscripted.")
        doc.save(source)

        formatter.post_process_docx(source, output, "Test Title", formatter.BUILTIN_PRESETS["academic-paper"])
        audit = formatter.audit_inline_citation_superscripts(output)
        require(audit["ok"], f"Expected inline citation audit to pass: {audit}")
        require(audit["inline_citation_count"] == 2, f"Expected two inline citations, got {audit}")

        ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
        with ZipFile(output) as archive:
            root = ET.fromstring(archive.read("word/document.xml"))

        citation_runs = {}
        reference_entry_superscripted = False
        interval_superscripted = False
        for para in root.findall(".//w:p", ns):
            para_text = "".join(t.text or "" for t in para.findall(".//w:t", ns))
            for run in para.findall("./w:r", ns):
                run_text = "".join(t.text or "" for t in run.findall(".//w:t", ns))
                if "[" not in run_text:
                    continue
                r_pr = run.find("./w:rPr", ns)
                vert = r_pr.find("./w:vertAlign", ns) if r_pr is not None else None
                value = vert.get("{%s}val" % ns["w"]) if vert is not None else None
                if run_text in {"[1-2]", "[3\uff0c4]"}:
                    citation_runs[run_text] = value
                if para_text.startswith("[1] Reference") and value == "superscript":
                    reference_entry_superscripted = True
                if "[0,1]" in run_text and value == "superscript":
                    interval_superscripted = True

        require(citation_runs == {"[1-2]": "superscript", "[3\uff0c4]": "superscript"}, f"Unexpected citation runs: {citation_runs}")
        require(not reference_entry_superscripted, "Reference-list entry was incorrectly superscripted.")
        require(not interval_superscripted, "Numeric interval [0,1] was incorrectly treated as a citation.")
    finally:
        shutil.rmtree(tmp_root, ignore_errors=True)


def english_support_smoke():
    formatter = load_formatter()
    require(
        formatter.infer_preset_name(Path("Safety Report.md"), "Safety Risk Report") == "english-report",
        "English report should infer english-report preset.",
    )
    require(
        formatter.infer_preset_name(Path("Research Paper.md"), "Knowledge Graph Paper") == "english-academic",
        "English paper should infer english-academic preset.",
    )
    for caption in ["Figure 1 Workflow", "Fig. 2: Result", "Table 1 Summary"]:
        require(formatter.looks_like_caption(caption), f"English caption was not recognized: {caption}")

    args = SimpleNamespace(final=None, output_dir=None, output_name=None, preset="auto", resolved_preset="english-report")
    final_path = formatter.final_docx_path(Path("Safety Report.md"), args)
    require(final_path.name == "Safety Report_formatted.docx", f"Unexpected English output name: {final_path.name}")


def main():
    validate_skill_metadata()
    validate_agent_portability_text()
    compile_scripts()
    install_target_resolution_smoke()
    citation_superscript_smoke()
    english_support_smoke()
    print("word-skill smoke tests passed")


if __name__ == "__main__":
    main()
